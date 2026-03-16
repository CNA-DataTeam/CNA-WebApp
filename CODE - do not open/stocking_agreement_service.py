"""
stocking_agreement_service.py

Helpers for building, rendering, and converting stocking agreement templates.
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

from docx import Document
from docx.table import _Row, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import Pt
from docxtpl import DocxTemplate
import pythoncom
from win32com import client as win32


APP_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = APP_DIR / "templates" / "stocking_agreements"
SOURCE_TEMPLATE_DIR = TEMPLATE_DIR / "source"
GENERAL_TEMPLATE = TEMPLATE_DIR / "general_resupply_template.docx"
CONSUMABLES_TEMPLATE = TEMPLATE_DIR / "consumables_template.docx"
LEGACY_GENERAL_SOURCE_TEMPLATE = SOURCE_TEMPLATE_DIR / "general_resupply_source.docx"
LEGACY_CONSUMABLES_SOURCE_TEMPLATE = SOURCE_TEMPLATE_DIR / "consumables_source.docx"
LEGACY_GENERAL_TEMPLATE = TEMPLATE_DIR / "general_resupply_template_v2.docx"
LEGACY_CONSUMABLES_TEMPLATE = TEMPLATE_DIR / "consumables_template_v2.docx"
LEGACY_TEMPLATE_TEMP_FILES = (
    TEMPLATE_DIR / "general_resupply_template.docx.tmp",
    TEMPLATE_DIR / "consumables_template.docx.tmp",
)
MIN_TEMPLATE_PRICING_ROWS = 6
DEFAULT_FONT_NAME = "Sofia Pro"
DEFAULT_FONT_SIZE_PT = 7

TEMPLATE_PATHS = {
    "general_resupply": GENERAL_TEMPLATE,
    "consumables": CONSUMABLES_TEMPLATE,
}


def ensure_templates_ready(force_rebuild: bool = False) -> dict[str, Path]:
    """Ensure one canonical template file exists for each agreement type."""
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

    _ensure_template_ready(
        target_path=GENERAL_TEMPLATE,
        legacy_render_path=LEGACY_GENERAL_TEMPLATE,
        legacy_source_path=LEGACY_GENERAL_SOURCE_TEMPLATE,
        builder=_build_general_resupply_template,
        force_rebuild=force_rebuild,
    )
    _ensure_template_ready(
        target_path=CONSUMABLES_TEMPLATE,
        legacy_render_path=LEGACY_CONSUMABLES_TEMPLATE,
        legacy_source_path=LEGACY_CONSUMABLES_SOURCE_TEMPLATE,
        builder=_build_consumables_template,
        force_rebuild=force_rebuild,
    )
    _cleanup_legacy_template_files()

    return dict(TEMPLATE_PATHS)


def render_agreement_documents(
    template_key: str,
    context: dict[str, Any],
    output_stem: str,
) -> dict[str, tuple[str, bytes]]:
    """Render the requested agreement template and return DOCX/PDF bytes."""
    templates = ensure_templates_ready()
    template_path = templates.get(template_key)
    if template_path is None:
        raise KeyError(f"Unsupported template key: {template_key}")
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    render_context = _prepare_render_context(template_key, context)
    with TemporaryDirectory(prefix="stocking_agreement_") as temp_dir:
        temp_root = Path(temp_dir)
        rendered_docx_path = temp_root / f"{output_stem}.docx"
        rendered_pdf_path = temp_root / f"{output_stem}.pdf"

        template = DocxTemplate(str(template_path))
        template.render(render_context, autoescape=False)
        template.save(str(rendered_docx_path))
        _populate_pricing_table(rendered_docx_path, template_key, context.get("pricing_rows", []))

        outputs: dict[str, tuple[str, bytes]] = {
            "docx": (f"{output_stem}.docx", rendered_docx_path.read_bytes())
        }

        try:
            pdf_bytes = _convert_docx_to_pdf_bytes(rendered_docx_path, rendered_pdf_path)
        except Exception as exc:
            outputs["pdf_error"] = (type(exc).__name__, str(exc))
        else:
            outputs["pdf"] = (f"{output_stem}.pdf", pdf_bytes)
        return outputs


def _ensure_template_ready(
    target_path: Path,
    legacy_render_path: Path,
    legacy_source_path: Path,
    builder,
    force_rebuild: bool,
) -> None:
    if force_rebuild and legacy_source_path.exists():
        builder(legacy_source_path, target_path)
        return

    if force_rebuild and legacy_render_path.exists():
        legacy_render_path.replace(target_path)
        return

    if target_path.exists() and not force_rebuild:
        return

    if legacy_render_path.exists():
        legacy_render_path.replace(target_path)
        return

    if legacy_source_path.exists():
        builder(legacy_source_path, target_path)
        return

    if target_path.exists():
        return

    raise FileNotFoundError(
        f"Agreement template not found: {target_path}. "
        "Keep one canonical template file per agreement type in the stocking agreements folder."
    )


def _cleanup_legacy_template_files() -> None:
    for path in (
        LEGACY_GENERAL_TEMPLATE,
        LEGACY_CONSUMABLES_TEMPLATE,
        *LEGACY_TEMPLATE_TEMP_FILES,
        LEGACY_GENERAL_SOURCE_TEMPLATE,
        LEGACY_CONSUMABLES_SOURCE_TEMPLATE,
    ):
        if path.exists():
            path.unlink()

    if SOURCE_TEMPLATE_DIR.exists() and not any(SOURCE_TEMPLATE_DIR.iterdir()):
        SOURCE_TEMPLATE_DIR.rmdir()


def _build_general_resupply_template(source_path: Path, output_path: Path) -> None:
    source_path = _require_source_template(source_path)
    doc = Document(str(source_path))

    details_table = doc.tables[0]
    _set_cell_text(
        details_table.rows[0].cells[1],
        "{{ project_name }}",
        donor_cell=details_table.rows[0].cells[0],
        fallback_size_pt=9,
    )
    _set_cell_text(
        details_table.rows[1].cells[1],
        "{{ account_name }}",
        donor_cell=details_table.rows[1].cells[0],
        fallback_size_pt=9,
    )

    _replace_first_paragraph_containing(
        doc,
        "This document serves to define the scope and parameters",
        (
            "This document serves to define the scope and parameters of the {{ project_name }} "
            "based on {{ client_name }} requirement and Clark National Accounts' capabilities. "
            "{{ executive_summary }}"
        ),
    )

    pricing_table = doc.tables[3]
    for idx in range(MIN_TEMPLATE_PRICING_ROWS):
        row = pricing_table.rows[idx + 1]
        _set_cell_text(row.cells[0], f"{{{{ pricing_rows[{idx}].description }}}}", fallback_size_pt=8)
        _set_cell_text(row.cells[1], f"{{{{ pricing_rows[{idx}].ea_sell }}}}", fallback_size_pt=8)
        _set_cell_text(row.cells[2], f"{{{{ pricing_rows[{idx}].qty }}}}", fallback_size_pt=8)
        _set_cell_text(row.cells[3], f"{{{{ pricing_rows[{idx}].extended_sell }}}}", fallback_size_pt=8)

    _set_cell_text(
        pricing_table.rows[7].cells[-1],
        "{{ item_subtotal }}",
        donor_cell=pricing_table.rows[7].cells[0],
        fallback_size_pt=8,
    )
    _set_cell_text(
        pricing_table.rows[8].cells[-1],
        "{{ freight }}",
        donor_cell=pricing_table.rows[8].cells[0],
        fallback_size_pt=8,
    )
    _set_cell_text(
        pricing_table.rows[9].cells[-1],
        "{{ tax }}",
        donor_cell=pricing_table.rows[9].cells[0],
        fallback_size_pt=8,
    )
    _set_cell_text(
        pricing_table.rows[10].cells[-1],
        "{{ order_total }}",
        donor_cell=pricing_table.rows[10].cells[0],
        fallback_size_pt=8,
    )

    project_details_table = doc.tables[5]
    _set_cell_text(
        project_details_table.rows[0].cells[1],
        "{{ shipping_method }}",
        donor_cell=project_details_table.rows[0].cells[0],
        fallback_size_pt=9,
    )
    _set_cell_text(
        project_details_table.rows[1].cells[1],
        "{{ payment_terms }}",
        donor_cell=project_details_table.rows[1].cells[0],
        fallback_size_pt=9,
    )
    _set_cell_text(
        project_details_table.rows[2].cells[1],
        "{{ invoicing }}",
        donor_cell=project_details_table.rows[2].cells[0],
        fallback_size_pt=9,
    )

    timeline_table = doc.tables[7]
    _set_cell_text(
        timeline_table.rows[0].cells[1],
        "{{ required_delivery_date }}",
        donor_cell=timeline_table.rows[0].cells[0],
        fallback_size_pt=9,
    )
    _set_cell_text(
        timeline_table.rows[1].cells[1],
        "{{ required_ship_date }}",
        donor_cell=timeline_table.rows[1].cells[0],
        fallback_size_pt=9,
    )

    misc_table = doc.tables[9]
    _set_cell_text(
        misc_table.rows[0].cells[1],
        "{{ documentation_requirements }}",
        donor_cell=misc_table.rows[0].cells[0],
        fallback_size_pt=9,
    )
    _set_cell_text(
        misc_table.rows[1].cells[1],
        "{{ addon_services_requested }}",
        donor_cell=misc_table.rows[1].cells[0],
        fallback_size_pt=9,
    )

    _replace_first_paragraph_containing(
        doc,
        "up to 21 days from the signed date of this agreement",
        (
            "3 CHANGES. Client shall have the right to make changes to project contents, "
            "packaging requirements, time and place of delivery and method of transportation "
            "{{ change_window_text }} unless otherwise outlined hereunder. If any such change "
            "causes an increase or decrease in the cost of, or the time required for the "
            "fulfillment of any part of the work under this contract, whether changed or not "
            "changed by any such order, Clark National Accounts reserves the right to make an "
            "equitable adjustment in the price, delivery schedule, or form of delivery, or all "
            "or any combination of the three, and this adjustment will be defined in a "
            "supplemental agreement that, once signed by both parties, will alter this agreement "
            "accordingly. Failure to agree to any adjustment shall be a dispute within the "
            "meaning of the Dispute clause of this contract."
        ),
    )

    _replace_first_paragraph_containing(
        doc,
        "termination charge consisting of ten percent",
        (
            "4.1 Termination for Convenience of Client. Client reserves the right to terminate "
            "this project or portion of the project hereof for its sole convenience. In the "
            "event of such termination, Clark National Accounts shall promptly stop all work "
            "hereunder, and shall promptly cause any of its suppliers and manufacturers to cease "
            "such work. Clark National Accounts shall be paid a termination charge "
            "{{ termination_charge_text }}."
        ),
    )

    _save_document(doc, output_path)


def _build_consumables_template(source_path: Path, output_path: Path) -> None:
    source_path = _require_source_template(source_path)
    doc = Document(str(source_path))

    details_table = doc.tables[0]
    _set_cell_text(
        details_table.rows[0].cells[1],
        "{{ project_name }}",
        donor_cell=details_table.rows[0].cells[0],
        fallback_size_pt=9,
    )
    _set_cell_text(
        details_table.rows[1].cells[1],
        "{{ account_name }}",
        donor_cell=details_table.rows[1].cells[0],
        fallback_size_pt=9,
    )

    _replace_first_paragraph_containing(
        doc,
        "This document serves to define the scope and parameters",
        (
            "This document serves to define the scope and parameters of the {{ project_name }} "
            "based on {{ account_name }} requirement and Clark National Accounts' capabilities."
        ),
    )

    _replace_first_paragraph_containing(
        doc,
        "FOR CONSOLIDATED ORDERS:",
        (
            "{% if order_type == 'Consolidated' %}FOR CONSOLIDATED ORDERS: Clark National "
            "Accounts (CNA) will be bringing in stock of {{ item_summary }} for {{ account_name }}. "
            "{{ purpose_details }} Once inventory lands, CNA will place a consolidated order, "
            "sending {{ cases_each }} CASES/EA to {{ location_count }} stores (totaling a final "
            "shipment of {{ total_units }} of {{ primary_item_number }}). Any remaining stock "
            "will be stored for a period of {{ storage_term_days }}-days following the receipt "
            "of the POs. During this time, the items will be merchandised and sold on the "
            "{{ account_name }} CNA site for applicable end users. For all inventory that CNA "
            "purchases specifically for {{ account_name }}, {{ account_name }} will be "
            "responsible for 100% of inventory that does not sell.{% endif %}"
        ),
    )

    _replace_first_paragraph_containing(
        doc,
        "FOR RESUPPLY ORDERS:",
        (
            "{% if order_type == 'Resupply' %}FOR RESUPPLY ORDERS: Clark National Accounts "
            "(CNA) will be bringing in stock of {{ item_summary }} for {{ account_name }}. "
            "{{ purpose_details }} CNA will store the inventory for a period of "
            "{{ storage_term_days }}-days following the receipt of the POs. During this time, "
            "the items will be merchandised and sold on the {{ account_name }} CNA site for "
            "applicable end users. For all inventory that CNA purchases specifically for "
            "{{ account_name }}, {{ account_name }} will be responsible for 100% of inventory "
            "that does not sell.{% endif %}"
        ),
    )

    _replace_first_paragraph_containing(
        doc,
        "At the end of the 90-day term",
        (
            "At the end of the {{ storage_term_days }}-day term, any inventory that remains will "
            "acquire a 5% storage fee per month based on the current sell price of the remaining "
            "stock. The storage fee will be charged to the {{ account_name }} corporate entity "
            "and standard payment of credit card payment is applicable in this circumstance. If "
            "inventory is deemed obsolete by {{ account_name }}, {{ account_name }} has the "
            "option to buy out the remaining stock at the current sell price. Per CNA's protocol, "
            "consumables may not be sold within 30-days of the product expiring. For any "
            "inventory on-hand, with an expiration date within 30-days, {{ account_name }} will "
            "have the option to purchase that inventory at the current sell price or can choose "
            "to have it disposed of at their expense."
        ),
    )

    pricing_table = doc.tables[3]
    while len(pricing_table.rows) < MIN_TEMPLATE_PRICING_ROWS + 1:
        pricing_table.add_row()
    for idx in range(MIN_TEMPLATE_PRICING_ROWS):
        row = pricing_table.rows[idx + 1]
        _set_cell_text(row.cells[0], f"{{{{ pricing_rows[{idx}].description }}}}", fallback_size_pt=8)
        _set_cell_text(row.cells[1], f"{{{{ pricing_rows[{idx}].ea_sell }}}}", fallback_size_pt=8)
        _set_cell_text(row.cells[2], f"{{{{ pricing_rows[{idx}].qty }}}}", fallback_size_pt=8)

    timeline_table = doc.tables[5]
    _set_cell_text(
        timeline_table.rows[0].cells[1],
        "{{ required_start_date }}",
        donor_cell=timeline_table.rows[0].cells[0],
        fallback_size_pt=9,
    )
    _set_cell_text(
        timeline_table.rows[1].cells[1],
        "{{ required_end_date }}",
        donor_cell=timeline_table.rows[1].cells[0],
        fallback_size_pt=9,
    )

    _replace_first_paragraph_containing(
        doc,
        "If stocking fees or remaining inventory need to be billed out",
        (
            "If stocking fees or remaining inventory need to be billed out, the following "
            "billing address will be used for invoicing under {{ billing_account_name }} CNA "
            "account:"
        ),
    )
    _replace_first_paragraph_containing(
        doc,
        "[ACCOUNT BILLING ADDRESS]",
        "{{ billing_address }}",
        fallback_size_pt=8,
    )

    _save_document(doc, output_path)


def _convert_docx_to_pdf_bytes(docx_path: Path, pdf_path: Path) -> bytes:
    word_app = None
    document = None
    pythoncom.CoInitialize()
    try:
        word_app = win32.DispatchEx("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = 0
        document = word_app.Documents.Open(str(docx_path.resolve()))
        document.ExportAsFixedFormat(
            OutputFileName=str(pdf_path.resolve()),
            ExportFormat=17,
        )
        document.Close(False)
        document = None
        word_app.Quit()
        word_app = None

        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF conversion did not create output: {pdf_path}")
        return pdf_path.read_bytes()
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word_app is not None:
            try:
                word_app.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def _prepare_render_context(template_key: str, context: dict[str, Any]) -> dict[str, Any]:
    render_context = dict(context)
    pricing_rows = context.get("pricing_rows", [])
    render_context["pricing_rows"] = _pad_pricing_rows_for_render(template_key, pricing_rows)
    return render_context


def _pad_pricing_rows_for_render(template_key: str, pricing_rows: list[dict[str, Any]]) -> list[dict[str, str]]:
    include_extended_sell = template_key == "general_resupply"
    rows: list[dict[str, str]] = []
    for item in pricing_rows:
        row = {
            "description": str(item.get("description", "") or ""),
            "ea_sell": str(item.get("ea_sell", "") or ""),
            "qty": str(item.get("qty", "") or ""),
        }
        if include_extended_sell:
            row["extended_sell"] = str(item.get("extended_sell", "") or "")
        rows.append(row)

    blank_row = {
        "description": "",
        "ea_sell": "",
        "qty": "",
    }
    if include_extended_sell:
        blank_row["extended_sell"] = ""

    while len(rows) < MIN_TEMPLATE_PRICING_ROWS:
        rows.append(dict(blank_row))
    return rows


def _populate_pricing_table(docx_path: Path, template_key: str, pricing_rows: list[dict[str, Any]]) -> None:
    doc = Document(str(docx_path))

    normalized_rows = _normalize_pricing_rows(pricing_rows, include_extended_sell=template_key == "general_resupply")
    if template_key == "general_resupply":
        table = doc.tables[3]
        _replace_table_rows(
            table=table,
            row_start=1,
            row_stop=7,
            rows=normalized_rows,
            column_keys=["description", "ea_sell", "qty", "extended_sell"],
        )
    elif template_key == "consumables":
        table = doc.tables[3]
        _replace_table_rows(
            table=table,
            row_start=1,
            row_stop=len(table.rows),
            rows=normalized_rows,
            column_keys=["description", "ea_sell", "qty"],
        )
    else:
        raise KeyError(f"Unsupported template key: {template_key}")

    _save_document(doc, docx_path)


def _normalize_pricing_rows(
    pricing_rows: list[dict[str, Any]],
    include_extended_sell: bool,
) -> list[dict[str, str]]:
    normalized_rows: list[dict[str, str]] = []
    for item in pricing_rows:
        row = {
            "description": str(item.get("description", "") or ""),
            "ea_sell": str(item.get("ea_sell", "") or ""),
            "qty": str(item.get("qty", "") or ""),
        }
        if include_extended_sell:
            row["extended_sell"] = str(item.get("extended_sell", "") or "")
        if any(value.strip() for value in row.values()):
            normalized_rows.append(row)

    if normalized_rows:
        return normalized_rows

    blank_row = {
        "description": "",
        "ea_sell": "",
        "qty": "",
    }
    if include_extended_sell:
        blank_row["extended_sell"] = ""
    return [blank_row]


def _replace_table_rows(
    table: Table,
    row_start: int,
    row_stop: int,
    rows: list[dict[str, str]],
    column_keys: list[str],
) -> None:
    placeholder_rows = list(table.rows[row_start:row_stop])
    if not placeholder_rows:
        raise ValueError("Could not locate placeholder rows in pricing table.")

    template_row_xml = deepcopy(placeholder_rows[0]._tr)
    table_xml = table._tbl
    insert_index = table_xml.index(placeholder_rows[0]._tr)

    for row in placeholder_rows:
        table_xml.remove(row._tr)

    for offset, values in enumerate(rows):
        new_row_xml = deepcopy(template_row_xml)
        table_xml.insert(insert_index + offset, new_row_xml)
        new_row = _Row(new_row_xml, table)
        for cell, key in zip(new_row.cells, column_keys):
            _set_cell_text(cell, str(values.get(key, "") or ""))


def _replace_first_paragraph_containing(
    doc: Document,
    needle: str,
    replacement: str,
    fallback_font_name: str = DEFAULT_FONT_NAME,
    fallback_size_pt: float = DEFAULT_FONT_SIZE_PT,
) -> None:
    needle_lower = needle.lower()
    for paragraph in doc.paragraphs:
        if needle_lower in paragraph.text.lower():
            _set_paragraph_text(
                paragraph,
                replacement,
                fallback_font_name=fallback_font_name,
                fallback_size_pt=fallback_size_pt,
            )
            return
    raise ValueError(f"Could not find paragraph containing: {needle}")


def _set_cell_text(
    cell,
    text: str,
    donor_cell=None,
    fallback_font_name: str = DEFAULT_FONT_NAME,
    fallback_size_pt: float = DEFAULT_FONT_SIZE_PT,
) -> None:
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    donor_run = _first_non_red_run_in_cell(cell)
    if donor_run is None and donor_cell is not None:
        donor_run = _first_non_red_run_in_cell(donor_cell)
    _set_paragraph_text(
        paragraph,
        text,
        donor_run=donor_run,
        fallback_font_name=fallback_font_name,
        fallback_size_pt=fallback_size_pt,
    )


def _set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    donor_run: Run | None = None,
    fallback_font_name: str = DEFAULT_FONT_NAME,
    fallback_size_pt: float = DEFAULT_FONT_SIZE_PT,
) -> None:
    if donor_run is None:
        donor_run = _first_non_red_run_in_paragraph(paragraph)

    _clear_paragraph_runs(paragraph)
    run = paragraph.add_run()
    _apply_run_format(
        run,
        donor_run=donor_run,
        fallback_font_name=fallback_font_name,
        fallback_size_pt=fallback_size_pt,
    )
    run.text = text


def _clear_paragraph_runs(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._element)


def _first_non_red_run_in_cell(cell) -> Run | None:
    for paragraph in cell.paragraphs:
        run = _first_non_red_run_in_paragraph(paragraph)
        if run is not None:
            return run
    return None


def _first_non_red_run_in_paragraph(paragraph: Paragraph) -> Run | None:
    fallback_run: Run | None = None
    for run in paragraph.runs:
        if fallback_run is None:
            fallback_run = run
        if not _is_red_run(run):
            return run
    return fallback_run


def _is_red_run(run: Run) -> bool:
    color = run.font.color.rgb if run.font.color is not None else None
    return str(color or "").upper() in {"FF0000", "C00000"}


def _apply_run_format(
    target_run: Run,
    donor_run: Run | None,
    fallback_font_name: str,
    fallback_size_pt: float,
) -> None:
    if donor_run is not None:
        target_run.style = donor_run.style
        if donor_run.font.name:
            target_run.font.name = donor_run.font.name
        else:
            target_run.font.name = fallback_font_name
        if donor_run.font.size:
            target_run.font.size = donor_run.font.size
        else:
            target_run.font.size = Pt(fallback_size_pt)
        if donor_run.bold is not None:
            target_run.bold = donor_run.bold
        if donor_run.italic is not None:
            target_run.italic = donor_run.italic
        if donor_run.underline is not None:
            target_run.underline = donor_run.underline
        if donor_run.font.color is not None and donor_run.font.color.rgb and not _is_red_run(donor_run):
            target_run.font.color.rgb = donor_run.font.color.rgb
        return

    target_run.font.name = fallback_font_name
    target_run.font.size = Pt(fallback_size_pt)


def _save_document(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temp_path = path.with_suffix(f"{path.suffix}.tmp")
    doc.save(str(temp_path))
    temp_path.replace(path)


def _require_source_template(path: Path) -> Path:
    if not path.exists():
        raise FileNotFoundError(
            f"Legacy source agreement template not found: {path}. "
            "Keep the canonical agreement template in the stocking agreements folder."
        )
    return path
